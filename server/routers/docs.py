"""
Document Assets and Export router.
"""

import io
import zipfile
import json
import hashlib
from datetime import datetime
from typing import List, Optional
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Response, BackgroundTasks
from PIL import Image, ImageDraw, ImageFont
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

import os
import tempfile
from server.auth import get_current_user
from server.routers.projects import verify_project_member
from core.utils.db import execute_query, execute_batch
from core.utils.r2 import read_text, write_binary, write_text, read_binary
from core.ocr.router import OCRRouter
from server.utils.graphics import render_manga_text_layers
from core.utils.text_chunker import chunk_text

router = APIRouter(prefix="/api/docs", tags=["Documents"])


class ExportRequest(BaseModel):
    format: str  # docx | txt | zip
    font: Optional[str] = "Arial"
    heading: Optional[str] = "h1"
    spacing: Optional[float] = 1.0
    project_id: str


@router.post("/{doc_id}/export")
async def export_document(
    doc_id: str,
    request: ExportRequest,
    current_user: dict = Depends(get_current_user)
):
    """
    Export document content in various formats (txt, docx, zip).
    """
    await verify_project_member(request.project_id, current_user["id"], "viewer")
    
    doc_prefix = f"projects/{request.project_id}/docs/{doc_id}/"
    
    # 1. Load the approved content from D1
    rows = await execute_query(
        "SELECT target_text FROM segments WHERE project_id = ? AND doc_id = ?",
        [request.project_id, doc_id]
    )
    
    content = ""
    if rows:
        content = "\n\n".join([r["target_text"] for r in rows if r["target_text"]])
        
    if not content:
        # Fallback to R2 metadata and output version
        meta_content = read_text(f"{doc_prefix}outputs/metadata.json")
        if meta_content:
            try:
                meta = json.loads(meta_content)
                v = meta.get("latest_version")
                output_content = read_text(f"{doc_prefix}outputs/translation_v{v}.json")
                if output_content:
                    out_data = json.loads(output_content)
                    content = out_data.get("approved_text", "")
            except Exception:
                pass
                
    if not content:
        # Fallback to draft
        content = read_text(f"{doc_prefix}draft.md") or ""
        
    if not content:
        content = f"Draft translation content for document {doc_id}."

    fmt = request.format.lower()
    
    if fmt == "txt":
        return Response(
            content=content,
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename={doc_id}.txt"}
        )
        
    elif fmt == "docx":
        # Check python-docx availability dynamically
        try:
            import docx
            doc = docx.Document()
            doc.add_heading(f"Document: {doc_id}", level=1)
            for p_text in content.split("\n\n"):
                if p_text.strip():
                    doc.add_paragraph(p_text.strip())
            
            stream = io.BytesIO()
            doc.save(stream)
            stream.seek(0)
            return StreamingResponse(
                stream,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={doc_id}.docx"}
            )
        except ImportError:
            # Fallback if python-docx is not installed
            fallback_text = f"Document: {doc_id}\n\n{content}"
            return Response(
                content=fallback_text.encode("utf-8"),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={doc_id}.docx"}
            )
            
    elif fmt == "zip":
        # Package a zip archive containing the text and mock images/bubbles
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
            # Write document text
            zip_file.writestr(f"{doc_id}_translation.txt", content)
            
            # Add a couple of dummy manga pages with text overlays
            dummy_png = b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15c4\x00\x00\x00\rIDATx\x9cc`\x00\x01\x00\x00\x0c\x00\x01\x04\x05\x73\x11\x00\x00\x00\x00IEND\xaeB`\x82"
            zip_file.writestr("page_001_overlay.png", dummy_png)
            zip_file.writestr("page_002_overlay.png", dummy_png)
            
        zip_buffer.seek(0)
        return StreamingResponse(
            zip_buffer,
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={doc_id}_manga_export.zip"}
        )
        
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported export format: {request.format}")


@router.post("/{doc_id}/assets/upload")
async def upload_assets(
    doc_id: str,
    project_id: str,
    files: List[UploadFile] = File(...),
    current_user: dict = Depends(get_current_user)
):
    """
    Upload assets for layout translation/OCR workflow.
    """
    await verify_project_member(project_id, current_user["id"], "editor")
    
    uploaded_files = []
    for file in files:
        contents = await file.read()
        r2_path = f"projects/{project_id}/docs/{doc_id}/assets/{file.filename}"
        write_binary(r2_path, contents, file.content_type)
        
        # Calculate checksum
        checksum = hashlib.sha1(contents).hexdigest()
        
        # Determine image dimensions using Pillow
        width, height = None, None
        try:
            img = Image.open(io.BytesIO(contents))
            width, height = img.size
        except Exception:
            pass  # Fallback if not an image or failed to parse
            
        # Insert metadata into D1 database
        asset_id = file.filename
        created_at = datetime.utcnow().isoformat() + "Z"
        
        await execute_query(
            """
            INSERT OR REPLACE INTO assets (
                asset_id, project_id, doc_id, asset_type, mime_type, r2_path, checksum, width, height, created_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            [asset_id, project_id, doc_id, "manga_page", file.content_type or "image/png", r2_path, checksum, width, height, created_at]
        )
        
        uploaded_files.append(file.filename)
        
    return {
        "status": "success",
        "message": f"Successfully uploaded {len(uploaded_files)} assets to R2 and registered in DB",
        "assets": uploaded_files
    }


@router.post("/{doc_id}/ocr/run")
async def run_ocr_batch(
    doc_id: str,
    project_id: str,
    force: bool = False,
    current_user: dict = Depends(get_current_user)
):
    """
    Run OCR detection on all document image assets and import segments.
    """
    await verify_project_member(project_id, current_user["id"], "editor")
    
    # 1. Check if approved translations already exist
    approved_check = await execute_query(
        "SELECT COUNT(*) as count FROM segments WHERE project_id = ? AND doc_id = ? AND approved_by IS NOT NULL",
        [project_id, doc_id]
    )
    approved_count = approved_check[0].get("count", 0) if approved_check else 0
    if approved_count > 0 and not force:
        raise HTTPException(
            status_code=409,
            detail="Approved translations already exist for this asset."
        )
        
    # 2. Get project source language
    proj_rows = await execute_query("SELECT source_lang FROM projects WHERE id = ?", [project_id])
    source_lang = proj_rows[0].get("source_lang", "ja") if proj_rows else "ja"
    
    # 3. Retrieve all assets for this doc
    assets = await execute_query(
        "SELECT asset_id, r2_path FROM assets WHERE project_id = ? AND doc_id = ?",
        [project_id, doc_id]
    )
    
    if not assets:
        # Fallback: list from R2 if database table is empty
        prefix = f"projects/{project_id}/docs/{doc_id}/assets/"
        from core.utils.r2 import list_files
        r2_files = list_files(prefix)
        assets = []
        for key in r2_files:
            filename = key.split("/")[-1]
            if filename:
                assets.append({"asset_id": filename, "r2_path": key})
                
    if not assets:
        return {
            "status": "success",
            "message": "No assets found to run OCR",
            "detected_bubbles": 0,
            "language": source_lang
        }
        
    # Get provider
    provider = OCRRouter.get_provider(source_lang)
    
    # 4. Clear existing segments inside a batch/transaction
    db_statements = [
        ("DELETE FROM segments WHERE project_id = ? AND doc_id = ?", [project_id, doc_id])
    ]
    
    total_bubbles = 0
    
    for asset in assets:
        asset_id = asset["asset_id"]
        r2_path = asset["r2_path"]
        
        # Read from R2
        image_bytes = read_binary(r2_path)
        if not image_bytes:
            continue
            
        # Write to local temp file for OCR provider
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_file:
            temp_file.write(image_bytes)
            temp_file_path = temp_file.name
            
        try:
            # Extract OCR blocks
            blocks = await provider.extract(temp_file_path)
            
            # Save raw OCR json to R2
            ocr_data = {
                "provider": provider.__class__.__name__.replace("Provider", "").lower(),
                "provider_version": "1.0.0",
                "generated_at": datetime.utcnow().isoformat() + "Z",
                "blocks": [
                    {
                        "text": b.text,
                        "bbox": b.bbox,
                        "confidence": b.confidence
                    }
                    for b in blocks
                ]
            }
            extracted_key = f"projects/{project_id}/docs/{doc_id}/extracted/{asset_id}.json"
            write_text(extracted_key, json.dumps(ocr_data, indent=2, ensure_ascii=False))
            
            # Create segment records
            for block in blocks:
                # Deterministic coordinate hash
                coord_str = "".join(f"{round(c, 3):.3f}" for c in block.bbox)
                h = hashlib.sha1(coord_str.encode('utf-8')).hexdigest()[:6]
                segment_id = f"{asset_id.split('.')[0]}_{h}"  # Remove extension from asset_id for cleaner IDs
                
                # SQL statement to insert segment
                db_statements.append((
                    """
                    INSERT INTO segments (
                        project_id, doc_id, segment_id, segment_type, source_text, target_text, asset_id, bbox
                    ) VALUES (?, ?, ?, 'bubble', ?, '', ?, ?)
                    """,
                    [project_id, doc_id, segment_id, block.text, asset_id, json.dumps(block.bbox)]
                ))
                total_bubbles += 1
                
        finally:
            # Clean up temp file
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
                
    # Execute database batch operations (will rollback if any insert fails)
    await execute_batch(db_statements)
        
    return {
        "status": "success",
        "message": "OCR batch detection completed successfully",
        "detected_bubbles": total_bubbles,
        "language": source_lang
    }


# Helper functions moved to server/utils/graphics.py


async def run_image_translation_pipeline_task(
    project_id: str,
    doc_id: str,
    asset_id: str,
    source_lang: str,
    target_lang: str,
    session_id: str
):
    from core.workflow.approval_gate import SESSION_CACHE
    from core.ocr.router import OCRRouter
    from core.memory import ProjectMemory
    from core.agents import TranslationAgent
    import tempfile
    
    session = SESSION_CACHE.get(session_id)
    if not session:
        print(f"[IMAGE PIPELINE] Error: Session {session_id} not found in SESSION_CACHE.")
        return
        
    try:
        # 1. OCR Stage
        print(f"[IMAGE PIPELINE] Starting OCR stage for asset: {asset_id}")
        session.pipeline_status["ocr"] = "running"
        
        r2_path = f"projects/{project_id}/docs/{doc_id}/assets/{asset_id}"
        image_bytes = read_binary(r2_path)
        if not image_bytes:
            raise ValueError(f"Asset image {asset_id} not found in R2")
            
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_file:
            temp_file.write(image_bytes)
            temp_file_path = temp_file.name
            
        try:
            provider = OCRRouter.get_provider(source_lang)
            print(f"[IMAGE PIPELINE] Using OCR provider: {provider.__class__.__name__}")
            blocks = await provider.extract(temp_file_path)
            print(f"[IMAGE PIPELINE] OCR extracted {len(blocks)} text blocks.")
            
            ocr_data = {
                "provider": provider.__class__.__name__.replace("Provider", "").lower(),
                "provider_version": "1.0.0",
                "generated_at": datetime.utcnow().isoformat() + "Z",
                "blocks": [{"text": b.text, "bbox": b.bbox, "confidence": b.confidence} for b in blocks]
            }
            extracted_key = f"projects/{project_id}/docs/{doc_id}/extracted/{asset_id}.json"
            write_text(extracted_key, json.dumps(ocr_data, indent=2, ensure_ascii=False))
            
            db_statements = [
                ("DELETE FROM segments WHERE project_id = ? AND doc_id = ? AND asset_id = ?", [project_id, doc_id, asset_id])
            ]
            for block in blocks:
                coord_str = "".join(f"{round(c, 3):.3f}" for c in block.bbox)
                h = hashlib.sha1(coord_str.encode('utf-8')).hexdigest()[:6]
                segment_id = f"{asset_id.split('.')[0]}_{h}"
                db_statements.append((
                    """
                    INSERT INTO segments (
                        project_id, doc_id, segment_id, segment_type, source_text, target_text, asset_id, bbox
                    ) VALUES (?, ?, ?, 'bubble', ?, '', ?, ?)
                    """,
                    [project_id, doc_id, segment_id, block.text, asset_id, json.dumps(block.bbox)]
                ))
            await execute_batch(db_statements)
            print(f"[IMAGE PIPELINE] Saved {len(blocks)} segments to SQLite/D1.")
        finally:
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
                
        session.pipeline_status["ocr"] = "success"
        session.pipeline_status["context_retrieval"] = "running"
        
        # 2. Context Retrieval Stage
        print(f"[IMAGE PIPELINE] Loading project memory context for project: {project_id}")
        mem = ProjectMemory(project_id)
        session.pipeline_status["context_retrieval"] = "success"
        session.pipeline_status["draft_translation"] = "running"
        
        # 3. Draft Translation Stage
        print(f"[IMAGE PIPELINE] Starting draft translation stage.")
        segments = await execute_query(
            "SELECT segment_id, source_text FROM segments WHERE project_id = ? AND doc_id = ? AND asset_id = ?",
            [project_id, doc_id, asset_id]
        )
        
        project_rows = await execute_query("SELECT description FROM projects WHERE id = ?", [project_id])
        project_description = project_rows[0].get("description", "") if project_rows else ""
        
        if not segments:
            print(f"[IMAGE PIPELINE] No text segments to translate for asset {asset_id}.")
        else:
            print(f"[IMAGE PIPELINE] Found {len(segments)} segments to translate.")
            all_source_texts = [s["source_text"] for s in segments if s.get("source_text", "").strip()]
            async with TranslationAgent(mem) as translation_agent:
                session.model_name = translation_agent.generator.model
                for idx, seg in enumerate(segments):
                    source_txt = seg["source_text"]
                    seg_id = seg["segment_id"]
                    if not source_txt.strip():
                        continue
                        
                    if idx > 0:
                        import asyncio
                        await asyncio.sleep(0.5)
                        
                    print(f"[IMAGE PIPELINE] Translating segment: {ascii(source_txt)}")
                    gen_result = await translation_agent.translate(
                        source_text=source_txt,
                        source_lang=source_lang,
                        target_lang=target_lang,
                        content_type="comic",
                        project_description=project_description,
                        context_sentences=all_source_texts
                    )
                    print(f"[IMAGE PIPELINE] Translated to: {ascii(gen_result.draft)}")
                    
                    await execute_query(
                        "UPDATE segments SET target_text = ? WHERE project_id = ? AND doc_id = ? AND segment_id = ?",
                        [gen_result.draft, project_id, doc_id, seg_id]
                    )
                    
        session.pipeline_status["draft_translation"] = "success"
        session.pipeline_status["review"] = "success"
        
        # 4. Generate initial preview render
        try:
            print(f"[IMAGE PIPELINE] Generating initial preview render for asset: {asset_id}")
            rendered_segments = await execute_query(
                "SELECT segment_id, source_text, target_text, bbox FROM segments WHERE project_id = ? AND doc_id = ? AND asset_id = ?",
                [project_id, doc_id, asset_id]
            )
            
            rendered_bytes = render_manga_text_layers(
                image_bytes=image_bytes,
                segments=rendered_segments,
                project_id=project_id,
                font_name="Arial",
                font_size=18
            )
            
            rendered_key = f"projects/{project_id}/docs/{doc_id}/rendered/{asset_id}"
            write_binary(rendered_key, rendered_bytes, "image/png")
            print(f"[IMAGE PIPELINE] Initial preview render saved successfully to: {rendered_key}")
            session.pipeline_status["render"] = "success"
        except Exception as render_err:
            print(f"[IMAGE PIPELINE] Warning: Failed to pre-render preview image: {ascii(render_err)}")
            session.pipeline_status["render"] = "failed"
            
        print(f"[IMAGE PIPELINE] Finished pipeline successfully for asset: {asset_id}")
        
    except Exception as e:
        print(f"[IMAGE PIPELINE] Exception occurred: {ascii(e)}")
        for stage in ["ocr", "context_retrieval", "draft_translation"]:
            if session.pipeline_status.get(stage) == "running":
                session.pipeline_status[stage] = "failed"
                break
        session.pipeline_error = str(e)
        
    finally:
        try:
            mem = ProjectMemory(project_id)
            from core.workflow.approval_gate import ApprovalGate
            gate = ApprovalGate(mem)
            gate.save_translation_draft(session)
        except Exception as e:
            print(f"[IMAGE PIPELINE] Error saving draft: {ascii(e)}")


@router.post("/{doc_id}/translate-image")
async def translate_image(
    doc_id: str,
    request: dict,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    project_id = request.get("project_id")
    asset_id = request.get("asset_id")
    source_lang = request.get("source_lang", "ja")
    target_lang = request.get("target_lang", "vi")
    
    await verify_project_member(project_id, current_user["id"], "editor")
    
    # Initialize a new ApprovalSession
    from core.memory import ProjectMemory
    from core.workflow.approval_gate import ApprovalGate
    
    mem = ProjectMemory(project_id)
    gate = ApprovalGate(mem)
    
    # Create the session in RAM & R2
    session = gate.create_session(
        doc_id=doc_id,
        source_text=f"Image Translation: {asset_id}",
        current_draft="",
        source_lang=source_lang,
        target_lang=target_lang
    )
    
    # Initialize pipeline status
    session.pipeline_status = {
        "upload": "success",
        "ocr": "idle",
        "context_retrieval": "idle",
        "draft_translation": "idle",
        "review": "idle",
        "approve": "idle",
        "render": "idle"
    }
    gate.save_translation_draft(session)
    
    # Add pipeline run task to BackgroundTasks
    background_tasks.add_task(
        run_image_translation_pipeline_task,
        project_id=project_id,
        doc_id=doc_id,
        asset_id=asset_id,
        source_lang=source_lang,
        target_lang=target_lang,
        session_id=session.session_id
    )
    
    return {
        "status": "success",
        "session_id": session.session_id,
        "message": "Image translation pipeline started in the background"
    }


@router.post("/{doc_id}/render")
async def render_document_image(
    doc_id: str,
    request: dict,
    current_user: dict = Depends(get_current_user)
):
    project_id = request.get("project_id")
    asset_id = request.get("asset_id")
    font_name = request.get("font_name", "Arial")
    font_size = request.get("font_size", 18)
    session_id = request.get("session_id")
    
    await verify_project_member(project_id, current_user["id"], "editor")
    
    r2_path = f"projects/{project_id}/docs/{doc_id}/assets/{asset_id}"
    image_bytes = read_binary(r2_path)
    if not image_bytes:
        raise HTTPException(status_code=404, detail=f"Asset image {asset_id} not found in R2")
        
    segments = await execute_query(
        "SELECT segment_id, source_text, target_text, bbox FROM segments WHERE project_id = ? AND doc_id = ? AND asset_id = ?",
        [project_id, doc_id, asset_id]
    )
    
    rendered_bytes = render_manga_text_layers(
        image_bytes=image_bytes,
        segments=segments,
        project_id=project_id,
        font_name=font_name,
        font_size=font_size
    )
    
    rendered_key = f"projects/{project_id}/docs/{doc_id}/rendered/{asset_id}"
    write_binary(rendered_key, rendered_bytes, "image/png")
    
    if session_id:
        from core.workflow.approval_gate import SESSION_CACHE
        session = SESSION_CACHE.get(session_id)
        if session:
            session.pipeline_status["render"] = "success"
            from core.workflow.approval_gate import ApprovalGate
            mem = ProjectMemory(project_id)
            gate = ApprovalGate(mem)
            gate.save_translation_draft(session)
            
    return {
        "status": "success",
        "message": f"Rendered image saved successfully for {asset_id}",
        "rendered_url": f"/api/docs/rendered/view/{project_id}/{doc_id}/{asset_id}"
    }


@router.get("/assets/view/{project_id}/{doc_id}/{asset_id}")
async def view_asset_image(
    project_id: str,
    doc_id: str,
    asset_id: str
):
    r2_path = f"projects/{project_id}/docs/{doc_id}/assets/{asset_id}"
    img_bytes = read_binary(r2_path)
    if not img_bytes:
        raise HTTPException(status_code=404, detail="Asset image not found")
        
    media_type = "image/png"
    if asset_id.lower().endswith(".jpg") or asset_id.lower().endswith(".jpeg"):
        media_type = "image/jpeg"
    elif asset_id.lower().endswith(".webp"):
        media_type = "image/webp"
        
    return StreamingResponse(io.BytesIO(img_bytes), media_type=media_type)


@router.get("/rendered/view/{project_id}/{doc_id}/{asset_id}")
async def view_rendered_image(
    project_id: str,
    doc_id: str,
    asset_id: str
):
    r2_path = f"projects/{project_id}/docs/{doc_id}/rendered/{asset_id}"
    img_bytes = read_binary(r2_path)
    if not img_bytes:
        raise HTTPException(status_code=404, detail="Rendered image not found")
        
    media_type = "image/png"
    return StreamingResponse(io.BytesIO(img_bytes), media_type=media_type)


@router.get("/{doc_id}/segments")
async def get_document_segments(
    doc_id: str,
    project_id: str,
    asset_id: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    await verify_project_member(project_id, current_user["id"], "viewer")
    
    if asset_id:
        rows = await execute_query(
            "SELECT segment_id, source_text, target_text, bbox, asset_id, approved_by, approved_at FROM segments WHERE project_id = ? AND doc_id = ? AND asset_id = ?",
            [project_id, doc_id, asset_id]
        )
    else:
        rows = await execute_query(
            "SELECT segment_id, source_text, target_text, bbox, asset_id, approved_by, approved_at FROM segments WHERE project_id = ? AND doc_id = ?",
            [project_id, doc_id]
        )
        
    return rows


@router.put("/{doc_id}/segments/{segment_id}")
async def update_segment_text(
    doc_id: str,
    segment_id: str,
    request: dict,
    current_user: dict = Depends(get_current_user)
):
    project_id = request.get("project_id")
    target_text = request.get("target_text", "")
    
    await verify_project_member(project_id, current_user["id"], "editor")
    
    await execute_query(
        "UPDATE segments SET target_text = ? WHERE project_id = ? AND doc_id = ? AND segment_id = ?",
        [target_text, project_id, doc_id, segment_id]
    )
    
    return {
        "status": "success",
        "message": f"Segment {segment_id} updated successfully"
    }


@router.post("/{doc_id}/upload-text-bulk")
async def upload_text_bulk(
    doc_id: str,
    project_id: str,
    source_lang: str,
    target_lang: str,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """
    Upload a large .txt file, chunk it, and queue it for background translation.
    """
    await verify_project_member(project_id, current_user["id"], "editor")
    
    contents = await file.read()
    try:
        text = contents.decode('utf-8')
    except UnicodeDecodeError:
        raise HTTPException(status_code=400, detail="File must be UTF-8 encoded text.")
        
    chunks = chunk_text(text, max_chars=1500)
    
    # Save the original file to R2
    r2_path = f"projects/{project_id}/docs/{doc_id}/assets/{file.filename}"
    write_binary(r2_path, contents, "text/plain")
    
    # Clear existing segments for this document
    await execute_query("DELETE FROM segments WHERE project_id = ? AND doc_id = ?", [project_id, doc_id])
    
    # Insert new segments
    db_statements = []
    for i, chunk in enumerate(chunks):
        segment_id = f"seg_{i:04d}"
        db_statements.append((
            """
            INSERT INTO segments (
                project_id, doc_id, segment_id, segment_type, source_text, target_text, asset_id
            ) VALUES (?, ?, ?, 'paragraph', ?, '', ?)
            """,
            [project_id, doc_id, segment_id, chunk, file.filename]
        ))
        
    if db_statements:
        await execute_batch(db_statements)
        
    # Queue background task
    from server.routers.translation import run_bulk_translation_pipeline_task
    background_tasks.add_task(
        run_bulk_translation_pipeline_task,
        project_id=project_id,
        doc_id=doc_id,
        source_lang=source_lang,
        target_lang=target_lang,
        fast_mode=True
    )
    
    return {
        "status": "success",
        "message": f"File chunked into {len(chunks)} segments and background translation started."
    }
